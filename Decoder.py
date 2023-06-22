import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import os

# Read data from Excel file
xls = pd.ExcelFile(r'inset path here')
excel_data_df = pd.read_excel(xls, sheet_name='MCQs', usecols=['Question', '(a) answer', '(b) answer', '(c) answer', '(d) answer', 'Image 1', 'Image 2 '])
df1 = pd.read_excel(xls, 'MCQs')

# Determine the correct answer
correct_answers = []
for index, row in df1.iterrows():
    found_answer = False
    for option in ['(a) solution', '(b) solution', '(c) solution', '(d) solution']:
        if row[option] == 1 or row[option] == 1.0:
            correct_answers.append(option)
            found_answer = True
            break
    if not found_answer:
        correct_answers.append(None)

# Append the correct answer column to the DataFrame
excel_data_df = excel_data_df.assign(Correct_Answer=correct_answers)

document = Document()

image_folder = r'insert path here'

for index, row in excel_data_df.iterrows():
    # Insert image 1
    image1_code = row['Image 1']
    if pd.notnull(image1_code):
        image1_path_jpg = os.path.join(image_folder, f'{image1_code}.jpg')
        image1_path_png = os.path.join(image_folder, f'{image1_code}.png')
        if os.path.exists(image1_path_jpg):
            try:
                document.add_picture(image1_path_jpg, width=Inches(3))
            except:
                # Save the image as PNG if JPEG format is not supported
                if os.path.exists(image1_path_png):
                    document.add_picture(image1_path_png, width=Inches(3))
        elif os.path.exists(image1_path_png):
            document.add_picture(image1_path_png, width=Inches(3))
    
    # Insert question
    question = row['Question']
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(f'Question: {question}\n').bold = True
    
    # Insert options
    options = ['(a) answer', '(b) answer', '(c) answer', '(d) answer']
    for option in options:
        answer = row[option]
        p.add_run(f'{option}: {answer}\n')
    
    # Insert correct answer
    correct_answer = row['Correct_Answer']
    p.add_run(f'Correct Answer: {correct_answer}\n\n').bold = True
    
    # Insert image 2
    image2_code = row['Image 2 ']
    if pd.notnull(image2_code):
        image2_path_jpg = os.path.join(image_folder, f'{image2_code}.jpg')
        image2_path_png = os.path.join(image_folder, f'{image2_code}.png')
        if os.path.exists(image2_path_jpg):
            try:
                document.add_picture(image2_path_jpg, width=Inches(3))
            except:
                # Save the image as PNG if JPEG format is not supported
                if os.path.exists(image2_path_png):
                    document.add_picture(image2_path_png, width=Inches(3))
        elif os.path.exists(image2_path_png):
            document.add_picture(image2_path_png, width=Inches(3))

# Save the Word document
document.save('insert name here.docx')
